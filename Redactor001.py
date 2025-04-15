import os
import win32com.client
import nltk
from nltk import word_tokenize, pos_tag, ne_chunk
import re
import fitz  # PyMuPDF
import json
import logging
import datetime
import matplotlib.pyplot as plt

class Redactor:

    def __init__(self):
        print("Redactor class initialized")

    def input_file(self, file_name):
        """
        Checks if the file exists in the 'output/' folder of the current working directory.
        If the file exists, returns the file path. Otherwise, returns None and -1.

        Parameters:
            file_name (str): The name of the file to check.

        Returns:
            tuple: (file_path, file_extension) if valid, else (None, -1).
        """
        # Get the path to the 'input/' folder in the current working directory
        output_folder = os.path.join(os.getcwd(), 'input')

        # Construct the full file path
        file_path = os.path.join(output_folder, file_name)

        # Check if the file exists
        if not os.path.exists(file_path):
            print(f"Error: File '{file_name}' does not exist in the 'output/' folder.")
            return None, -1

        # Get the file extension
        _, file_extension = os.path.splitext(file_name)
        file_extension = file_extension.lower()

        return file_path, file_extension

    def __extract_text_txt(self, file_path):
        """
        Extracts text from a plain text file.

        Args:
            file_path (str): The path to the file from which to extract text.

        Returns:
            str: The extracted text.
        """
        # Extract text from a plain text file
        print ("Extracting text from a .txt or .md file")
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def __extract_text_pdf(self, file_path):
        """
        Extracts text from a PDF file.

        Args:
            file_path (str): The path to the PDF file from which to extract text.

        Returns:
            str: The extracted text.
        """
        # Extract text from a PDF file
        print ("Extracting text from a .pdf file")
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = ''
            for page in reader.pages:
                text += page.extract_text()
            return text
        except ImportError:
            return "Current Version only supports text pdfs, we are improving it to support images and scanned pdfs. Please install PyPDF2 using 'pip install PyPDF2'."

    def __extract_text_doc(self, file_path):
        """
        Extracts text from a Word document.

        Args:
            file_path (str): The path to the Word document from which to extract text.

        Returns:
            str: The extracted text.
        """
        # Extract text from a Word document
        try:
            # Initialize Word application
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Open the .doc file
            doc = word.Documents.Open(file_path)

            # Extract content by iterating through paragraphs
            content = []
            for paragraph in doc.Paragraphs:
                content.append(paragraph.Range.Text.strip())

            # Close the document and Word application
            doc.Close()
            word.Quit()

            # Join paragraphs with newlines
            return "\n".join(content)
        except Exception as e:
            print(f"Error: {e}")
            return None

    def __extract_text_docx(self, file_path):
        """
        Extracts text from a Word document.

        Args:
            file_path (str): The path to the Word document from which to extract text.

        Returns:
            str: The extracted text.
        """
        # Extract text from a Word document
        print ("Extracting text from a .docx file")
        try:
            from docx import Document
            doc = Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except ImportError:
            return "Current Version only supports text Word Docx, we are improving it to support images and scanned Docx. Install it using 'pip install python-docx'."

    def extract_text(self, file_path, file_extension):
        """
        Extracts text from a given file.

        Args:
            file_path (str): The path to the file from which to extract text.

        Returns:
            str: The extracted text, or an error message if the file type is unsupported.
        """
    
        # Check the file extension
        if (file_extension == '.txt') or (file_extension == '.md'):
            return self.__extract_text_txt(file_path)
            
        elif (file_extension =='.pdf'):
            return self.__extract_text_pdf(file_path)

        elif file_extension == '.docx':
            return self.__extract_text_docx(file_path)

        elif file_extension == '.doc':
            return self.__extract_text_doc(file_path)

        else:
            return "Unsupported file type. Please provide a .txt, .pdf, or .docx file."

    def obfuscate_words(self, doc_text, user_choice, user_options):
        """
        Obfuscates words in the document text based on user choice.

        Parameters:
            doc_text (str): The text of the document to be obfuscated.
            user_choice (str): The user's choice for obfuscation.

        Returns:
            list: A list of target words to be removed from the document.
        """
        # Check if user_choice is a valid option
        if user_choice.lower() == "custom":
            return self.__obfuscate_custom(doc_text, user_options['custom'])
        elif user_choice.lower() == "gradation":
            return self.__obfuscate_gradation(doc_text, user_options['gradation'])
        else:
            print("Invalid choice. Please select 'custom' or 'gradation'.")
            return []
    
    def __obfuscate_gradation(self, doc_text, level):
        """
        Obfuscates the document text based on the gradation fields provided.

        Parameters:
            doc_text (str): The text of the document to be obfuscated.
            fields (list): A list of words or phrases to be obfuscated.

        Returns:
            list: A list of target words to be removed from the document.
        """
        if level not in [1, 2, 3, 4]:
            print("Invalid level. Please select a level between 1 and 4.")
            return []
        tag_dir = {
            1: ["NN", "NNS", "NNP", "NNPS"],
            2: ["JJ", "JJR", "JJS"],
            3: ["RB", "RBR", "RBS"],
            4: ["VB", "VBD", "VBG", "VBN", "VBP", "VBZ"],
        }

        return self.__obfuscate_custom(doc_text, tag_dir[level])

    def __obfuscate_custom(self, doc_text, fields):
        """
        Obfuscates the document text based on the custom fields provided.

        Parameters:
            doc_text (str): The text of the document to be obfuscated.
            fields (list): A list of words or phrases to be obfuscated.

        Returns:
            list: A list of target words to be removed from the document.
        """
        pos = {
            "coordinating conjunction": "CC",
            "cardinal number": "CD",
            "determiner": "DT",
            "existential there": "EX",
            "foreign word": "FW",
            "preposition or subordinating conjunction": "IN",
            "adjective": "JJ",
            "adjective, comparative": "JJR",
            "adjective, superlative": "JJS",
            "list item marker": "LS",
            "modal": "MD",
            "noun, singular or mass": "NN",
            "noun, plural": "NNS",
            "proper noun, singular": "NNP",
            "proper noun, plural": "NNPS",
            "predeterminer": "PDT",
            "possessive ending": "POS",
            "personal pronoun": "PRP",
            "possessive pronoun": "PRP$",
            "adverb": "RB",
            "adverb, comparative": "RBR",
            "adverb, superlative": "RBS",
            "particle": "RP",
            "symbol": "SYM",
            "to": "TO",
            "interjection": "UH",
            "verb, base form": "VB",
            "verb, past tense": "VBD",
            "verb, gerund or present participle": "VBG",
            "verb, past participle": "VBN",
            "verb, non-3rd person singular present": "VBP",
            "verb, 3rd person singular present": "VBZ",
            "wh-determiner": "WDT",
            "wh-pronoun": "WP",
            "possessive wh-pronoun": "WP$",
            "wh-adverb": "WRB",
        }

        '''
        1.	CC	Coordinating conjunction
        2.	CD	Cardinal number
        3.	DT	Determiner
        4.	EX	Existential there
        5.	FW	Foreign word
        6.	IN	Preposition or subordinating conjunction
        7.	JJ	Adjective
        8.	JJR	Adjective, comparative
        9.	JJS	Adjective, superlative
        10.	LS	List item marker
        11.	MD	Modal
        12.	NN	Noun, singular or mass
        13.	NNS	Noun, plural
        14.	NNP	Proper noun, singular
        15.	NNPS	Proper noun, plural
        16.	PDT	Predeterminer
        17.	POS	Possessive ending
        18.	PRP	Personal pronoun
        19.	PRP$	Possessive pronoun
        20.	RB	Adverb
        21.	RBR	Adverb, comparative
        22.	RBS	Adverb, superlative
        23.	RP	Particle
        24.	SYM	Symbol
        25.	TO	to
        26.	UH	Interjection
        27.	VB	Verb, base form
        28.	VBD	Verb, past tense
        29.	VBG	Verb, gerund or present participle
        30.	VBN	Verb, past participle
        31.	VBP	Verb, non-3rd person singular present
        32.	VBZ	Verb, 3rd person singular present
        33.	WDT	Wh-determiner
        34.	WP	Wh-pronoun
        35.	WP$	Possessive wh-pronoun
        36.	WRB	Wh-adverb
        '''

        # Check if fields is a valid list
        if not isinstance(fields, list) or not all(isinstance(field, str) for field in fields):
            return []  # Return an empty list if fields are invalid

        # fields_tag = [pos[field.lower()] for field in fields if field in pos]
        fields_tag = [field for field in fields if field in pos.values()]
        # print ("Fields Tag:", fields_tag)

        # Tokenize the document text
        tokens = word_tokenize(doc_text)
        # print ("Tokens:", tokens)
        pos_tags = pos_tag(tokens)
        # print ("POS Tags:", pos_tags)
        # named_entities = ne_chunk(pos_tags)
        # print ("Named Entities:", named_entities)


        # Initialize the target words list
        target_words = []

        # Iterate through the fields and check if they exist in the tokens
        # for field in fields:
            # for chunk in named_entities:
            #     if hasattr(chunk, 'tag') and chunk.tag() == 'NN':
            #         # print ("Label: ", chunk)
            #         target_words.append(" ".join(c[0] for c in chunk))
        for word in pos_tags:
            if word[1] in fields_tag and len(word[0]) > 2:
                target_words.append(word[0])
        return target_words

    def __obfuscate_reconstruct_txt(self, file_path, target_words):
        """
        Replace all occurrences of a word in a text file and save the modified content to a new file.

        Parameters:
            file_path (str): Path to the input text file.
            target_word (str): The word to be replaced

        Returns:
            None
        """
        replacement_word = '[ REDACTED ]'
        output_file = 'output//' + os.path.basename(file_path)
        try:
            # Open the input file and read its content
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Replace all occurrences of the target word
                
            # Create a regex pattern to match any of the target words
            pattern = r'\b(' + '|'.join(map(re.escape, target_words)) + r')\b'
            
            # Replace all matches with the replacement word
            modified_content = re.sub(pattern, replacement_word, content)

            # Write the modified content to the output file
            with open(output_file, 'w', encoding='utf-8') as file:
                file.write(modified_content)

            print(f"Replaced '{target_words}' with '{replacement_word}' in {file_path} and saved to {output_file}.")
        except FileNotFoundError:
            print(f"Error: The file {file_path} was not found.")
        except Exception as e:
            print(f"An error occurred: {e}")

    def __obfuscate_reconstruct_pdf(self, file_path, target_words):
        """
        Replace words in a PDF while preserving its format and avoiding overlaps.

        Parameters:
            input_pdf (str): Path to the input PDF file.
            output_pdf (str): Path to save the modified PDF file.
            target_words (list): A list of words to be replaced with asterisks.

        Returns:
            None
        """
        # Open the PDF
        pdf_document = fitz.open(file_path)

        # Iterate through each page
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]

            # Search for target words and replace them
            for target_word in target_words:
                areas = page.search_for(target_word)  # Find all occurrences of the word
                for area in areas:
                    # Redact the word by overlaying a white rectangle
                    page.draw_rect(area, color=(1, 1, 1), fill=(1, 1, 1))  # White rectangle to clear the text
                    # Insert replacement text (e.g., asterisks) in the same position
                    replacement = "*" * len(target_word)
                    page.insert_text(area.tl, replacement, fontsize=10, color=(0, 0, 0))  # Black text

        # Save the modified PDF
        output_pdf = 'output//' + os.path.basename(file_path)
        # Ensure the output directory exists
        pdf_document.save(output_pdf)
        pdf_document.close()

    def __obfuscate_reconstruct_docx(self, file_path, target_words):
        """
        Replace words in a Word document while preserving its format and avoiding overlaps.

        Parameters:
            input_doc (str): Path to the input Word document.
            output_doc (str): Path to save the modified Word document.
            target_words (list): A list of words to be replaced with asterisks.

        Returns:
            None
        """
        # Open the Word document
        try:
            from docx import Document
            doc = Document(file_path)

            # Iterate through paragraphs and replace target words
            for paragraph in doc.paragraphs:
                for target_word in target_words:
                    if target_word in paragraph.text:
                        paragraph.text = paragraph.text.replace(target_word, "*" * len(target_word))

            # Save the modified document
            output_doc = 'output//' + os.path.basename(file_path)
            doc.save(output_doc)
        except Exception as e:
            print(f"Error: {e}")

    def __obfuscate_reconstruct_doc(self, file_path, target_words):
        """
        Replace words in a Word document while preserving its format and avoiding overlaps.

        Parameters:
            input_doc (str): Path to the input Word document.
            output_doc (str): Path to save the modified Word document.
            target_words (list): A list of words to be replaced with asterisks.

        Returns:
            None
        """
        # Open the Word document
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Open the .doc file
            doc = word.Documents.Open(file_path)

            # Iterate through paragraphs and replace target words
            for paragraph in doc.Paragraphs:
                for target_word in target_words:
                    paragraph.Range.Find.Execute(FindText=target_word, ReplaceWith="*" * len(target_word), Replace=2)

            # Save the modified document
            output_doc = 'output//' + os.path.basename(file_path)
            doc.SaveAs(output_doc)
            doc.Close()
            word.Quit()
        except Exception as e:
            print(f"Error: {e}")
    
    def __obfuscate_text(self, file_path, target_words):
        """
        Replace all occurrences of a word in a text file and save the modified content to a new file.

        Parameters:
            file_path (str): Path to the input text file.
            target_word (str): The word to be replaced

        Returns:
            None
        """
        replacement_word = '[ REDACTED ]'
        output_file = 'output//' + os.path.basename(file_path)
        try:
            # Open the input file and read its content
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Replace all occurrences of the target word
                
            # Create a regex pattern to match any of the target words
            pattern = r'\b(' + '|'.join(map(re.escape, target_words)) + r')\b'
            
            # Replace all matches with the replacement word
            modified_content = re.sub(pattern, replacement_word, content)

            # Write the modified content to the output file
            with open(output_file, 'w', encoding='utf-8') as file:
                file.write(modified_content)

            print(f"Replaced '{target_words}' with '{replacement_word}' in {file_path} and saved to {output_file}.")
        except FileNotFoundError:
            print(f"Error: The file {file_path} was not found.")
        except Exception as e:
            print(f"An error occurred: {e}")

    def obfuscate(self, file_path, target_words, output_choice):
        """
        Obfuscates the document text based on user choice.

        Parameters:
            doc_text (str): The text of the document to be obfuscated.
            user_choice (str): The user's choice for obfuscation.

        Returns:
            list: A list of target words to be removed from the document.
        """
        # Check if user_choice is a valid option
        if output_choice.lower() == "txt":
            self.__obfuscate_reconstruct_txt(file_path, target_words)
        elif output_choice.lower() == "pdf":
            self.__obfuscate_reconstruct_pdf(file_path, target_words)
        elif output_choice.lower() == "docx":
            self.__obfuscate_reconstruct_docx(file_path, target_words)
        elif output_choice.lower() == "doc":
            self.__obfuscate_reconstruct_doc(file_path, target_words)
        else:
            self.__obfuscate_text(file_path, target_words)

    def generate_analysis_and_report(self, file_path, doc_text, target_words, replacement_word):
        """
        Generates an analysis and report for the obfuscation process.

        Parameters:
            file_path (str): Path to the input file.
            target_words (list): List of words that were obfuscated.
            replacement_word (str): The word or pattern used for replacement.
            output_file (str): Path to the obfuscated output file.

        Returns:
            dict: Analysis data.
        """
        try:
            
            original_content = doc_text

            # Calculate statistics
            total_words = len(original_content.split())
            obfuscated_count = sum(original_content.count(word) for word in target_words)
            obfuscation_percentage = (obfuscated_count / total_words) * 100 if total_words > 0 else 0

            # Prepare analysis data
            analysis_data = {
                "file_name": os.path.basename(file_path),
                "file_type": os.path.splitext(file_path)[1],
                "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "total_words": total_words,
                "obfuscated_count": obfuscated_count,
                "obfuscation_percentage": round(obfuscation_percentage, 2),
                "target_words": target_words,
                "replacement_word": replacement_word
            }

            # Save the analysis as a JSON report
            report_file = os.path.splitext(file_name)[0] + "_analysis.json"
            with open(report_file, 'w', encoding='utf-8') as report:
                json.dump(analysis_data, report, indent=4)

            print(f"Analysis and report saved to {report_file}.")
            return analysis_data

        except Exception as e:
            print(f"Error generating analysis and report: {e}")
            return None
            
    def generate_graphs(self, analysis_data):
        """
        Generates graphs based on the analysis data.

        Parameters:
            analysis_data (dict): The analysis data containing statistics like total words, obfuscated count, etc.

        Returns:
            None
        """
        try:
            # Extract data from analysis
            total_words = analysis_data["total_words"]
            obfuscated_count = analysis_data["obfuscated_count"]
            obfuscation_percentage = analysis_data["obfuscation_percentage"]

            # Bar Chart: Total Words vs Obfuscated Words
            plt.figure(figsize=(8, 6))
            plt.bar(["Total Words", "Obfuscated Words"], [total_words, obfuscated_count], color=["blue", "red"])
            plt.title("Total Words vs Obfuscated Words")
            plt.ylabel("Count")
            plt.savefig("output/word_analysis_bar_chart.png")
            plt.show()

            # Pie Chart: Obfuscation Percentage
            plt.figure(figsize=(8, 6))
            plt.pie(
                [obfuscated_count, total_words - obfuscated_count],
                labels=["Obfuscated", "Non-Obfuscated"],
                autopct="%1.1f%%",
                colors=["red", "green"],
                startangle=90,
            )
            plt.title("Obfuscation Percentage")
            plt.savefig("output/obfuscation_percentage_pie_chart.png")
            plt.show()

            print("Graphs generated and saved in the 'output/' folder.")

        except Exception as e:
            print(f"Error generating graphs: {e}")

readactor = Redactor()
with open('input.json', 'r', encoding='utf-8') as file:
    input_parameters = json.load(file)
file_name = input_parameters['file_name']
user_choice = input_parameters['user_choice']
user_options = input_parameters['user_options']
output_choice = input_parameters['output_choice']

file_path, file_extension = readactor.input_file(file_name)
if file_path is None:
    exit(-1)
doc_text = readactor.extract_text(file_path, file_extension)
if doc_text is None:
    exit(-1)
target_words = readactor.obfuscate_words(doc_text, user_choice, user_options)
if target_words is None:
    exit(-1)
readactor.obfuscate(file_path, target_words, output_choice)
print("Obfuscation completed successfully.")
analysis_data = readactor.generate_analysis_and_report(file_path, doc_text, target_words, '[ REDACTED ]')
readactor.generate_graphs(analysis_data)